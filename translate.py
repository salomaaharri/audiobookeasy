#!/usr/bin/env python3
"""
translate.py — Translate a DOCX chapter from English to Finnish using Claude Sonnet 4.6.

Preserves paragraph structure (Heading 1/2, body) so the output can feed into docx2mp3.py.
"""

import argparse
import os
import re
import sys
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from anthropic import Anthropic


SYSTEM_PROMPT = """You are a master literary translator working from English to Finnish.

Your task is to translate fantasy / sci-fi / literary novel chapters at the highest possible literary quality, producing prose that reads as if originally written in Finnish.

Translation principles:
- Preserve the original literary register, voice, mood, and pacing.
- Use natural, idiomatic Finnish — avoid translationese, awkward calques, and over-literal renderings.
- Keep English proper names (characters, places, invented terms, fictional technologies) EXACTLY as in the source.
- Render dialogue naturally — match the speaker's register (formal/casual, terse/expansive) in Finnish.
- For invented or technical English terms with no Finnish equivalent: keep the English term.
- Preserve emphasis (italics suggested by tone, exclamations, ellipses) where it carries meaning.

Input format:
The user will send a numbered list of paragraphs in this exact format:

[1] <text of paragraph 1>

[2] <text of paragraph 2>

[3] <text of paragraph 3>

Output format:
Output EXCLUSIVELY the Finnish translation in the same numbered format:

[1] <Finnish translation of paragraph 1>

[2] <Finnish translation of paragraph 2>

[3] <Finnish translation of paragraph 3>

Critical rules:
- Translate EVERY input paragraph — no merging, splitting, omitting, or reordering.
- Use the EXACT same numbering as input ([1], [2], [3], ...).
- NO commentary, NO preamble, NO explanations, NO trailing notes — only the numbered translated paragraphs.
- If a paragraph is a chapter heading or short title (e.g. "Chapter 1: The Sky Falls Again"), translate it but keep it as a single short paragraph."""


def load_dotenv(*candidates: Path) -> None:
    """Minimal .env loader (no python-dotenv dependency)."""
    for path in candidates:
        if not path.is_file():
            continue
        for raw in path.read_text(encoding="utf-8", errors="ignore").splitlines():
            line = raw.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            if line.startswith("export "):
                line = line[len("export "):]
            key, _, val = line.partition("=")
            key = key.strip()
            val = val.strip()
            if len(val) >= 2 and val[0] == val[-1] and val[0] in ("'", '"'):
                val = val[1:-1]
            if key and key not in os.environ:
                os.environ[key] = val
        return


def read_paragraphs(src: Path) -> List[Tuple[str, str]]:
    """Read non-empty paragraphs as (style_name, text) tuples."""
    doc = Document(str(src))
    items: List[Tuple[str, str]] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            style = getattr(p.style, "name", "") or "Normal"
            items.append((style, text))
    return items


def format_for_translation(items: List[Tuple[str, str]]) -> str:
    """Format paragraphs as a numbered list for the translator."""
    lines = []
    for i, (_, text) in enumerate(items, 1):
        lines.append(f"[{i}] {text}")
    return "\n\n".join(lines)


def parse_translated(text: str, n_expected: int) -> List[Optional[str]]:
    """Parse '[N] <text>' chunks back into a list of N translations (None for missing)."""
    pattern = re.compile(r"^\[(\d+)\]\s*([\s\S]*?)(?=^\[\d+\]|\Z)", re.MULTILINE)
    result: List[Optional[str]] = [None] * n_expected
    for m in pattern.finditer(text):
        idx = int(m.group(1)) - 1
        if 0 <= idx < n_expected:
            result[idx] = m.group(2).strip()
    return result


def write_docx(items: List[Tuple[str, str]], translations: List[Optional[str]], out_path: Path) -> None:
    """Write a new DOCX preserving the original paragraph styles, with translated text."""
    doc = Document()
    available_styles = {s.name for s in doc.styles}
    for (style, orig), trans in zip(items, translations):
        text = trans if trans is not None else orig  # fall back to original if translation missing
        p = doc.add_paragraph()
        if style in available_styles:
            try:
                p.style = style
            except Exception:
                pass
        p.add_run(text)
    doc.save(str(out_path))


def translate_chapter(client: Anthropic, src: Path, out_path: Path, model: str, max_tokens: int) -> bool:
    items = read_paragraphs(src)
    if not items:
        print(f"  ! no paragraphs in {src}", file=sys.stderr)
        return False

    prompt = format_for_translation(items)
    n = len(items)
    total_chars = sum(len(t) for _, t in items)
    print(f"  paragraphs: {n} | chars: {total_chars:,}")

    user_msg = (
        f"Translate this English novel chapter into Finnish. "
        f"There are {n} paragraphs. Output the Finnish translation in the same numbered format.\n\n"
        f"{prompt}"
    )

    with client.messages.stream(
        model=model,
        max_tokens=max_tokens,
        thinking={"type": "adaptive"},
        output_config={"effort": "medium"},
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_msg}],
    ) as stream:
        response = stream.get_final_message()

    text = next((b.text for b in response.content if b.type == "text"), "")
    translations = parse_translated(text, n)

    n_found = sum(1 for t in translations if t)
    print(f"  translated: {n_found}/{n}")

    if response.stop_reason == "max_tokens":
        print(f"  ! stop_reason=max_tokens — output may be truncated", file=sys.stderr)
    if n_found < n:
        missing = [i + 1 for i, t in enumerate(translations) if not t]
        print(f"  ! missing paragraphs: {missing}", file=sys.stderr)

    write_docx(items, translations, out_path)

    u = response.usage
    print(f"  usage: input={u.input_tokens} | output={u.output_tokens}"
          + (f" | cache_read={u.cache_read_input_tokens}" if getattr(u, "cache_read_input_tokens", 0) else ""))
    return n_found >= n  # success only if every paragraph came back


def main() -> None:
    parser = argparse.ArgumentParser(description="Translate a DOCX from English to Finnish via Claude Sonnet 4.6.")
    parser.add_argument("source", type=Path, help="Input .docx file")
    parser.add_argument("--out", type=Path, required=True, help="Output .docx path")
    parser.add_argument("--model", default="claude-sonnet-4-6", help="Anthropic model ID (default: claude-sonnet-4-6)")
    parser.add_argument("--max-tokens", type=int, default=24000, help="max_tokens for the API call (default: 24000)")
    args = parser.parse_args()

    script_dir = Path(__file__).resolve().parent
    load_dotenv(
        script_dir / ".env.local",
        script_dir / ".env",
        Path.cwd() / ".env.local",
        Path.cwd() / ".env",
    )

    if not os.environ.get("ANTHROPIC_API_KEY"):
        sys.exit("✖ ANTHROPIC_API_KEY not set. Add it to .env.local or export it.")

    if not args.source.is_file():
        sys.exit(f"✖ Source not found: {args.source}")

    args.out.parent.mkdir(parents=True, exist_ok=True)

    client = Anthropic()
    print(f"→ {args.source.name}")
    ok = translate_chapter(client, args.source, args.out, args.model, args.max_tokens)
    if ok:
        print(f"✓ {args.out}")
    else:
        print(f"✗ {args.out} (partial — review before use)")
        sys.exit(2)


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""
    docx2mp3.py — Turn DOCX/TXT into an MP3 audiobook (per-chapter files + combined)
    using Microsoft Edge TTS (edge-tts), with clean chapter detection and robust chunking.

    ──────────────────────────────────────────────────────────────────────────────
    🔹 What this script does
      • Reads a .docx or .txt manuscript
      • Detects chapters from DOCX headings (Heading 1/2 / Otsikko 1/2)
        or heuristically from lines like "Luku", "Osa", "Chapter"
      • Splits chapter text into TTS-safe chunks (~2200 chars) with sentence-aware logic
      • Synthesizes each chunk to MP3 via edge-tts (text mode; no SSML needed)
      • Exports:
          - One MP3 per chapter with ID3 tags (album/artist/title)
          - One combined MP3 (all chapters with small gaps between)

    🔧 Requirements
      pip install edge-tts python-docx pydub tqdm
      FFmpeg must be installed and on your PATH (pydub uses it for MP3 I/O)

    💡 Examples
      doc2mp3.py book.docx --outdir out --album "My Audiobook" --author "Your Name"
      doc2mp3.py book.docx --voice fi-FI-SelmaNeural --rate -5 --volume +3
      doc2mp3.py notes.txt --no-per-chapter  # only export a single combined file

    Author: Harri J. Salomaa
    """

import argparse
import asyncio
import os
import re
import shutil
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Protocol, Tuple

from tqdm import tqdm
from pydub import AudioSegment

def existing_doc(path_str: str) -> Path:
    p = Path(path_str).expanduser()
    if not p.exists():
        raise argparse.ArgumentTypeError(f"Source not found: {p}")
    if not p.is_file():
        raise argparse.ArgumentTypeError(f"Source is not a file: {p}")
    if p.suffix.lower() not in {".docx", ".txt"}:
        raise argparse.ArgumentTypeError("Source must be .docx or .txt")
    return p.resolve()

def writable_dir(path_str: str) -> Path:
    d = Path(path_str).expanduser()
    try:
        d.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        raise argparse.ArgumentTypeError(f"Cannot create output dir '{d}': {e}")
    # quick writability test
    try:
        (d / ".write_test").touch(exist_ok=True)
        (d / ".write_test").unlink(missing_ok=True)
    except Exception as e:
        raise argparse.ArgumentTypeError(f"Output dir not writable '{d}': {e}")
    return d.resolve()

def require_ffmpeg():
    if shutil.which("ffmpeg") is None:
        sys.exit("✖ FFmpeg not found on PATH. Install it and try again.")


def load_dotenv(*candidates: Path) -> None:
    """
    Tiny .env loader (no python-dotenv dependency). Reads KEY=VALUE lines from the
    first existing file in `candidates` and sets them in os.environ without
    overwriting variables already present in the real environment.
    """
    for path in candidates:
        if not path.is_file():
            continue
        for raw in path.read_text(encoding="utf-8", errors="ignore").splitlines():
            line = raw.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            if line.startswith("export "):
                line = line[len("export "):]
            key, _, val = line.partition("=")
            key = key.strip()
            val = val.strip()
            # Strip surrounding single or double quotes
            if len(val) >= 2 and val[0] == val[-1] and val[0] in ("'", '"'):
                val = val[1:-1]
            if key and key not in os.environ:
                os.environ[key] = val
        return  # only load the first file that exists

# ===============================
# Small helpers
# ===============================
def slugify(name: str, max_len: int = 80) -> str:
    # Remove illegal filesystem chars (Windows/macOS/Linux safe)
    name = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", name)
    # Collapse whitespace to underscores
    name = re.sub(r"\s+", "_", name.strip())
    # Trim leading/trailing dots/underscores
    name = name.strip("._")
    # Cap length so chapter-title-derived filenames stay below filesystem limits
    if len(name) > max_len:
        name = name[:max_len].rstrip("._")
    return name or "chapter"

def derive_title_prefix(src: Path, chapters: List["Chapter"]) -> str:
    # 1) DOCX document properties title
    if src.suffix.lower() == ".docx":
        try:
            from docx import Document
            t = (Document(str(src)).core_properties.title or "").strip()
            if t:
                return slugify(t)
        except Exception:
            pass
    # 2) First heading if present (and not Untitled)
    if chapters and chapters[0].title and chapters[0].title != "Untitled":
        return slugify(chapters[0].title)
    # 3) Fallback to input file name
    return slugify(src.stem)

def ensure_percent(val: str) -> str:
    """
    Normalize numeric or dB inputs to percentages required by edge-tts:
        - '-5'   → '-5%'
        - '+3dB' → '+3%'
        - '+3'   → '+3%'
    """
    s = str(val).strip()
    if s.lower().endswith("db"):
        s = s[:-2]
    return s if s.endswith("%") else f"{s}%"

# ===============================
# Chapter reading (DOCX / TXT)
# ===============================
@dataclass
class Chapter:
    title: str
    text: str

def read_docx_chapters(path: Path) -> List[Chapter]:
    """
    Read a DOCX file, detect chapters by Heading 1/2 (English: 'Heading 1/2',
    Finnish: 'Otsikko 1/2'). If no headings exist, fall back to heuristic detection.
    """
    from docx import Document

    doc = Document(str(path))
    # Collect (style_name, paragraph_text)
    items: List[Tuple[str, str]] = []
    for p in doc.paragraphs:
        style = getattr(p.style, "name", "") or ""
        text = (p.text or "").strip()
        if text:
            items.append((style, text))

    chapters: List[Chapter] = []
    current_title: Optional[str] = None
    buffer: List[str] = []

    def flush():
        nonlocal current_title, buffer
        content = "\n\n".join(buffer).strip()
        if content:
            chapters.append(Chapter(title=current_title or "Untitled", text=content))
        current_title = None
        buffer = []

    for style, text in items:
        is_heading = any(h in style for h in ("Heading 1", "Heading 2", "Otsikko 1", "Otsikko 2"))
        # Require a number after Luku/Chapter/Osa — these are common Finnish/English words that
        # otherwise produce false positives on body paragraphs (e.g. "Osa merkeistä oli...").
        is_luku_like = bool(re.match(r"^(Luku|Chapter|Osa)\s+\d", text, flags=re.IGNORECASE))
        if is_heading or is_luku_like:
            if buffer:
                flush()
            current_title = text
        else:
            buffer.append(text)
    if buffer:
        flush()

    # Fallback: whole document as one chapter if detection failed
    if not chapters:
        full = "\n\n".join(t for _, t in items)
        chapters = [Chapter(title="Book", text=re.sub(r"\n{3,}", "\n\n", full).strip())]
    return chapters

def read_txt_chapters(path: Path) -> List[Chapter]:
    """
    Heuristic chapter detection for plain text:
        - Splits when a line starts with 'Luku', 'Osa', or 'Chapter' (optionally with numbers).
    If nothing detected, returns a single 'Book' chapter.
    """
    txt = path.read_text(encoding="utf-8", errors="ignore")
    lines = [l.rstrip() for l in txt.splitlines()]

    chapters: List[Chapter] = []
    current_title: Optional[str] = None
    buffer: List[str] = []

    def flush():
        nonlocal current_title, buffer
        content = "\n".join(buffer).strip()
        if content:
            chapters.append(Chapter(title=current_title or f"Chapter {len(chapters)+1}", text=content))
        current_title = None
        buffer = []

    for line in lines:
        if re.match(r"^(Luku\s+\d+|Chapter\s+\d+|Osa\s+\d+|Luku\b|Chapter\b|Osa\b)", line, flags=re.IGNORECASE):
            if buffer:
                flush()
            current_title = line.strip()
        else:
            buffer.append(line)
    if buffer:
        flush()

    if not chapters:
        chapters = [Chapter(title="Book", text=re.sub(r"\n{3,}", "\n\n", txt).strip())]
    return chapters

def read_chapters(path: Path) -> List[Chapter]:
    """Dispatch based on extension."""
    if path.suffix.lower() == ".docx":
        return read_docx_chapters(path)
    return read_txt_chapters(path)

# ===============================
# Chunking for TTS
# ===============================
def split_chunks(text: str, max_chars: int = 2200) -> List[str]:
    """
    Split text into TTS-friendly chunks.
    Prefers paragraph -> sentence boundaries, then hard-splits if necessary.
    """
    parts: List[str] = []
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if len(para) <= max_chars:
            parts.append(para)
            continue

        sentences = re.split(r"(?<=[.!?…])\s+", para)
        buf = ""
        for s in sentences:
            if not s:
                continue
            if len(buf) + 1 + len(s) <= max_chars:
                buf = f"{buf} {s}".strip() if buf else s
            else:
                if buf:
                    parts.append(buf)
                if len(s) <= max_chars:
                    buf = s
                else:
                    for i in range(0, len(s), max_chars):
                        parts.append(s[i : i + max_chars])
                    buf = ""
        if buf:
            parts.append(buf)
    return parts or [text]

# ===============================
# TTS backends
# ===============================
@dataclass
class SynthOptions:
    voice: str
    rate: str = "-5%"
    volume: str = "+0%"
    instructions: Optional[str] = None      # openai gpt-4o-mini-tts
    stability: Optional[float] = None       # elevenlabs
    similarity: Optional[float] = None      # elevenlabs
    style: Optional[float] = None           # elevenlabs
    model: Optional[str] = None             # engine model override


class TTSBackend(Protocol):
    name: str
    default_voice: str
    default_model: str
    async def synth(self, text: str, outfile: Path, opts: SynthOptions) -> None: ...


class EdgeBackend:
    name = "edge"
    default_voice = "fi-FI-SelmaNeural"
    default_model = ""

    async def synth(self, text: str, outfile: Path, opts: SynthOptions) -> None:
        import edge_tts
        communicate = edge_tts.Communicate(
            text=text, voice=opts.voice, rate=opts.rate, volume=opts.volume
        )
        await communicate.save(str(outfile))


def _rate_pct_to_speed(rate_pct: str) -> float:
    """Map Edge-style percentage rate (e.g. '-5%') to OpenAI speed (0.25–4.0)."""
    s = str(rate_pct).strip().rstrip("%")
    try:
        return max(0.25, min(4.0, 1.0 + float(s) / 100.0))
    except ValueError:
        return 1.0


class OpenAIBackend:
    name = "openai"
    default_voice = "alloy"
    default_model = "gpt-4o-mini-tts"

    async def synth(self, text: str, outfile: Path, opts: SynthOptions) -> None:
        try:
            from openai import AsyncOpenAI
        except ImportError:
            sys.exit("✖ openai SDK not installed. Run: pip install 'openai>=1.40.0'")

        client = AsyncOpenAI()  # reads OPENAI_API_KEY
        kwargs = {
            "model": opts.model or self.default_model,
            "voice": opts.voice,
            "input": text,
            "response_format": "mp3",
            "speed": _rate_pct_to_speed(opts.rate),
        }
        if opts.instructions:
            kwargs["instructions"] = opts.instructions
        async with client.audio.speech.with_streaming_response.create(**kwargs) as resp:
            await resp.stream_to_file(str(outfile))


class ElevenLabsBackend:
    name = "elevenlabs"
    default_voice = ""  # required — voice IDs are account-specific
    default_model = "eleven_multilingual_v2"

    async def synth(self, text: str, outfile: Path, opts: SynthOptions) -> None:
        if not opts.voice:
            sys.exit("✖ ElevenLabs requires --voice <voice_id>. Find IDs in your ElevenLabs dashboard.")
        try:
            from elevenlabs.client import AsyncElevenLabs
        except ImportError:
            sys.exit("✖ elevenlabs SDK not installed. Run: pip install 'elevenlabs>=1.0.0'")

        client = AsyncElevenLabs()  # reads ELEVENLABS_API_KEY
        audio = client.text_to_speech.convert(
            voice_id=opts.voice,
            model_id=opts.model or self.default_model,
            text=text,
            output_format="mp3_44100_128",
            voice_settings={
                "stability": 0.5 if opts.stability is None else opts.stability,
                "similarity_boost": 0.75 if opts.similarity is None else opts.similarity,
                "style": 0.0 if opts.style is None else opts.style,
                "use_speaker_boost": True,
            },
        )
        with open(outfile, "wb") as f:
            async for chunk in audio:
                if chunk:
                    f.write(chunk)


_BACKENDS = {
    "edge": EdgeBackend,
    "openai": OpenAIBackend,
    "elevenlabs": ElevenLabsBackend,
}


def get_backend(name: str) -> TTSBackend:
    if name not in _BACKENDS:
        sys.exit(f"✖ Unknown engine: {name}. Choose from {sorted(_BACKENDS)}.")
    return _BACKENDS[name]()

async def synth_chapter(chapter: Chapter, tmpdir: Path, backend: TTSBackend, opts: SynthOptions) -> AudioSegment:
    """
    Synthesize a whole chapter by chunking it and concatenating the resulting MP3 segments.
    """
    parts = split_chunks(chapter.text, max_chars=2200)
    pause = AudioSegment.silent(duration=800)  # small gap between chunks
    assembled = AudioSegment.empty()
    for idx, part in enumerate(parts):
        part_file = tmpdir / f"{slugify(chapter.title)}_{idx:04d}.mp3"
        # Retry per-chunk: OpenAI's streaming TTS occasionally drops connections mid-body
        # (httpx.RemoteProtocolError / httpcore.ReadError). One bad chunk shouldn't kill
        # a 30-minute chapter, so back off and retry the chunk itself.
        last_exc: Optional[BaseException] = None
        for attempt in range(1, 6):
            try:
                await backend.synth(part, part_file, opts)
                last_exc = None
                break
            except Exception as e:
                last_exc = e
                if attempt == 5:
                    break
                delay = min(2 ** attempt, 30)
                print(f"  ! chunk {idx} attempt {attempt} failed ({type(e).__name__}: {e}); retrying in {delay}s", flush=True)
                await asyncio.sleep(delay)
        if last_exc is not None:
            raise last_exc
        assembled += AudioSegment.from_file(part_file, format="mp3")
        if idx < len(parts) - 1:
            assembled += pause
    return assembled

# ===============================
# Pipeline
# ===============================
async def build(
    src: Path,
    outdir: Path,
    album: str,
    author: str,
    backend: TTSBackend,
    opts: SynthOptions,
    per_chapter: bool = True,
    combined_name: str = "book_combined.mp3",
    chapter_gap_ms: int = 1200,
    bitrate: str = "192k",
    prefix: Optional[str] = None,
) -> Path:
    """
    Full end-to-end pipeline:
      1) Read chapters
      2) Synthesize each chapter
      3) Export per-chapter MP3s (optional)
      4) Export combined MP3 with gaps between chapters
    """
    chapters = read_chapters(src)
    if chapters and chapters[0].title == "Untitled":
        chapters[0].title = "Esipuhe"

    # Compute filename prefix (explicit → auto-derived)
    auto_prefix = derive_title_prefix(src, chapters)
    effective_prefix = slugify(prefix) if prefix else auto_prefix

    outdir.mkdir(parents=True, exist_ok=True)

    combined = AudioSegment.empty()
    with tempfile.TemporaryDirectory() as td:
        tdir = Path(td)
        for idx, ch in enumerate(tqdm(chapters, desc="Chapters", unit="chapter")):
            seg = await synth_chapter(ch, tdir, backend, opts)

            if per_chapter:
                fname = f"{effective_prefix}_{idx+1:02d}_{slugify(ch.title)}.mp3"
                fpath = outdir / fname
                seg.export(
                    fpath, format="mp3", bitrate=bitrate,
                    tags={"album": album, "artist": author, "title": ch.title}
                )

            combined += seg
            if idx < len(chapters) - 1:
                combined += AudioSegment.silent(duration=chapter_gap_ms)

    # Prefix the combined file too
    combined_path = outdir / f"{effective_prefix}_{combined_name}"
    combined.export(
        combined_path, format="mp3", bitrate=bitrate,
        tags={"album": album, "artist": author, "title": src.stem}
    )
    return combined_path

# ===============================
# CLI
# ===============================
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="DOCX/TXT → per-chapter MP3 files + combined audiobook (Edge / OpenAI / ElevenLabs TTS)."
    )
    parser.add_argument("source", type=existing_doc, help="Input .docx or .txt")
    parser.add_argument("--outdir", type=writable_dir, default=Path("output_mp3"), help="Output folder (default: output_mp3)")
    parser.add_argument("--album", type=str, default="Audiobook", help="Album tag for MP3 metadata")
    parser.add_argument("--author", type=str, default="Unknown Author", help="Artist/Author tag for MP3 metadata")

    # Engine selection
    parser.add_argument("--engine", choices=sorted(_BACKENDS), default="edge",
                        help="TTS engine (default: edge)")
    parser.add_argument("--voice", default=None,
                        help="Voice (engine-specific). Defaults: edge=fi-FI-SelmaNeural, openai=alloy, elevenlabs=(required)")
    parser.add_argument("--model", default=None,
                        help="Engine model override (openai: gpt-4o-mini-tts, elevenlabs: eleven_multilingual_v2)")

    # Edge / OpenAI shared
    parser.add_argument("--rate", default="-5%", help="Speech rate (edge: -5/-5%%; openai: mapped to speed 0.25–4.0)")
    parser.add_argument("--volume", default="+0%", help="Volume (edge only; accepts +3, +3%%, +3dB)")

    # OpenAI only
    parser.add_argument("--instructions", default=None,
                        help="OpenAI gpt-4o-mini-tts style prompt (e.g. 'Warm, contemplative narration; slow on emotional beats')")

    # ElevenLabs only
    parser.add_argument("--stability", type=float, default=None, help="ElevenLabs stability (0.0–1.0, default 0.5)")
    parser.add_argument("--similarity", type=float, default=None, help="ElevenLabs similarity_boost (0.0–1.0, default 0.75)")
    parser.add_argument("--style", type=float, default=None, help="ElevenLabs style (0.0–1.0, default 0.0)")

    # Output options
    parser.add_argument("--no-per-chapter", action="store_true", help="Skip per-chapter export; only write combined file")
    parser.add_argument("--combined-name", default="book_combined.mp3", help="Filename for combined audiobook")
    parser.add_argument("--chapter-gap-ms", type=int, default=1200, help="Silence between chapters in combined file (ms)")
    parser.add_argument("--bitrate", default="192k", help="MP3 bitrate (128k–320k)")
    parser.add_argument("--prefix", default=None, help="Filename prefix (defaults to DOCX title → first heading → input stem)")

    args = parser.parse_args()

    # Load API keys from .env.local / .env in the script's directory (CWD as fallback)
    script_dir = Path(__file__).resolve().parent
    load_dotenv(
        script_dir / ".env.local",
        script_dir / ".env",
        Path.cwd() / ".env.local",
        Path.cwd() / ".env",
    )

    # Pre-flight checks
    require_ffmpeg()
    if args.engine == "openai" and not os.environ.get("OPENAI_API_KEY"):
        sys.exit("✖ OPENAI_API_KEY not set. export OPENAI_API_KEY=... and retry.")
    if args.engine == "elevenlabs" and not os.environ.get("ELEVENLABS_API_KEY"):
        sys.exit("✖ ELEVENLABS_API_KEY not set. export ELEVENLABS_API_KEY=... and retry.")

    # Normalize percent/dB inputs (edge-tts expects percent strings)
    args.rate = ensure_percent(args.rate)
    args.volume = ensure_percent(args.volume)

    # Build backend + options
    backend = get_backend(args.engine)
    voice = args.voice if args.voice is not None else backend.default_voice
    opts = SynthOptions(
        voice=voice,
        rate=args.rate,
        volume=args.volume,
        instructions=args.instructions,
        stability=args.stability,
        similarity=args.similarity,
        style=args.style,
        model=args.model,
    )

    asyncio.run(
        build(
            src=args.source,
            outdir=args.outdir,
            album=args.album,
            author=args.author,
            backend=backend,
            opts=opts,
            per_chapter=not args.no_per_chapter,
            combined_name=args.combined_name,
            chapter_gap_ms=args.chapter_gap_ms,
            bitrate=args.bitrate,
            prefix=args.prefix,
        )
    )
